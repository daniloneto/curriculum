from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.parser import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.shared import Inches
import os
import json
import sys
import glob
import argparse
from templates import TemplateManager

# Configurar os argumentos de linha de comando
parser = argparse.ArgumentParser(description='Gerar currículo em formato DOCX.')
parser.add_argument('language', nargs='?', help='Código do idioma (ex: pt, en, es)')
parser.add_argument('--template', '-t', help='Nome do template a ser usado', default='docx')
parser.add_argument('--json-file', help='Caminho para um arquivo JSON personalizado', default=None)
args = parser.parse_args()

# Função para listar idiomas disponíveis
def get_available_languages():
    # Procurar todos os arquivos JSON que seguem o padrão curriculo_XX.json
    json_files = glob.glob('curriculo_*.json')
    languages = {}
    
    for file in json_files:
        # Extrair o código do idioma do nome do arquivo (curriculo_XX.json -> XX)
        lang_code = file.replace('curriculo_', '').replace('.json', '')
        
        # Carregar o arquivo para obter o nome do idioma na própria língua
        try:
            with open(file, 'r', encoding='utf-8') as f:
                data = json.load(f)
                
                # Verificar se o arquivo tem a estrutura esperada
                if 'languageName' in data:
                    lang_name = data['languageName']
                else:
                    # Fallback para casos onde o nome do idioma não está definido
                    lang_name = lang_code.upper()
                
                languages[lang_code] = {
                    'name': lang_name,
                    'file': file
                }
        except:
            # Se houver erro ao carregar ou analisar, pular este arquivo
            pass
    
    return languages

# Determinar o idioma a ser usado
available_languages = get_available_languages()

# Default para português se disponível, caso contrário usa o primeiro idioma disponível
default_lang = 'pt' if 'pt' in available_languages else list(available_languages.keys())[0] if available_languages else None

# Verificar qual idioma usar com base nos argumentos de linha de comando
selected_lang = default_lang
if args.language:
    lang_arg = args.language.lower()
    if lang_arg in available_languages:
        selected_lang = lang_arg

# Se não houver idiomas disponíveis, terminar o programa
if not selected_lang and not args.json_file:
    print("Erro: Não foram encontrados arquivos de idioma válidos.")
    sys.exit(1)

# Carregar o arquivo JSON
if args.json_file:
    print(f"Usando arquivo JSON personalizado: {args.json_file}")
    # Usar o arquivo JSON personalizado
    try:
        with open(args.json_file, 'r', encoding='utf-8') as file:
            data = json.load(file)
    except Exception as e:
        print(f"Erro ao carregar arquivo JSON personalizado: {str(e)}")
        sys.exit(1)
else:
    # Carregar o arquivo JSON do idioma selecionado
    json_file = available_languages[selected_lang]['file']
    with open(json_file, 'r', encoding='utf-8') as file:
        data = json.load(file)

# Carregar o template
template_manager = TemplateManager()
template_name = args.template

try:
    template = template_manager.get_template(template_name)
    print(f"Usando template: {template_name}")
except ValueError as e:
    print(f"Erro ao carregar template: {str(e)}")
    print(f"Templates disponíveis: {', '.join(template_manager.list_templates())}")
    print("Usando o template padrão 'docx'.")
    template = template_manager.get_template('docx')

# Novo documento
doc = template.create_document()

# Função genérica para obter valores do JSON de maneira padronizada
def get_field(data, primary_key, fallback_key=None, additional_fallbacks=None):
    if primary_key in data:
        return data[primary_key]
    elif fallback_key and fallback_key in data:
        return data[fallback_key]
    
    # Verificar chaves adicionais específicas para alguns campos
    if additional_fallbacks:
        for key in additional_fallbacks:
            if key in data:
                return data[key]
    
    return None

# Função genérica para obter título de seção
def get_section_title(section_data, title_keys=['titulo', 'title', 'titre', 'titel']):
    for key in title_keys:
        if key in section_data:
            return section_data[key]
    return "N/A"  # Fallback

# Função genérica para obter conteúdo de seção
def get_section_content(section_data, content_keys=['conteudo', 'content', 'contenido', 'inhalt']):
    for key in content_keys:
        if key in section_data:
            return section_data[key]
    return ""  # Fallback

# Função genérica para obter lista de itens
def get_section_list(section_data, list_keys=['lista', 'list', 'liste', 'lista']):
    for key in list_keys:
        if key in section_data:
            return section_data[key]
    return []  # Fallback

# Função genérica para obter lista de empregos/experiências
def get_jobs(section_data, jobs_keys=['empregos', 'jobs', 'empleos', 'emplois']):
    for key in jobs_keys:
        if key in section_data:
            return section_data[key]
    return []  # Fallback

# Extrair dados básicos do JSON
# Estrutura padronizada para todas as línguas
nome = get_field(data, 'nome', 'name', ['nombre'])
email = data['email']
telefone = get_field(data, 'telefone', 'phone')
linkedin = data['linkedin']

# Determinar qual é a chave principal para seções
secoes_key = None
for key in ['secoes', 'sections', 'secciones', 'sektionen']:
    if key in data:
        secoes_key = key
        break

# Se não encontrarmos a chave das seções, não podemos continuar
if not secoes_key:
    print("Erro: Formato de arquivo JSON inválido. A chave de seções não foi encontrada.")
    sys.exit(1)

secoes = data[secoes_key]

# Montando o currículo visual
template.add_title(doc, nome, email, telefone, linkedin)

# Resumo Profissional - procurar por várias possíveis chaves
resume_section = None
for key in ['resumoProfissional', 'professionalSummary', 'resumenProfesional', 'resumentProfessionnel']:
    if key in secoes:
        resume_section = secoes[key]
        break

if resume_section:
    template.add_section_title(doc, get_section_title(resume_section))
    doc.add_paragraph(get_section_content(resume_section))

# Experiência Profissional - procurar por várias possíveis chaves
experience_section = None
for key in ['experienciaProfissional', 'workExperience', 'experienciaLaboral', 'experienceProfessionnelle']:
    if key in secoes:
        experience_section = secoes[key]
        break

if experience_section:
    template.add_section_title(doc, get_section_title(experience_section))
    
    # Obter lista de empregos
    jobs = get_jobs(experience_section)
    
    # Adicionar empregos
    for job in jobs:
        position = get_field(job, 'cargo', 'position')
        if position:
            doc.add_paragraph(position, style='List Bullet')
        
        period = get_field(job, 'periodo', 'period')
        if period:
            doc.add_paragraph(period)
        
        # Obter descrição - procurar por várias possíveis chaves
        description_items = []
        for key in ['descricao', 'description', 'descripcion']:
            if key in job:
                description_items = job[key]
                break
        
        # Montar descrição
        descricao = ""
        for item in description_items:
            descricao += f"- {item}\n"
        doc.add_paragraph(descricao)

# Habilidades Técnicas - procurar por várias possíveis chaves
template.add_page_break(doc)
skills_section = None
for key in ['habilidadesTecnicas', 'technicalSkills', 'habilidadesTecnicas', 'competencesTechniques']:
    if key in secoes:
        skills_section = secoes[key]
        break

if skills_section:
    template.add_section_title(doc, get_section_title(skills_section))
    
    # Obter lista de habilidades
    skills = []
    for key in ['habilidades', 'skills', 'habilidades', 'competences']:
        if key in skills_section:
            skills = skills_section[key]
            break
    
    for skill in skills:
        skill_name = get_field(skill, 'nome', 'name')
        skill_level = get_field(skill, 'nivel', 'level')
        if skill_name and skill_level:
            template.add_skill_bar(doc, skill_name, skill_level)

# Certificações - procurar por várias possíveis chaves
certifications_section = None
for key in ['certificacoes', 'certifications', 'certificaciones', 'certifications']:
    if key in secoes:
        certifications_section = secoes[key]
        break

if certifications_section:
    template.add_section_title(doc, get_section_title(certifications_section))
    certifications = get_section_list(certifications_section)
    
    for cert in certifications:
        p = doc.add_paragraph()
        p.add_run("🏅 ").bold = True
        p.add_run(cert)

# Educação - procurar por várias possíveis chaves
education_section = None
for key in ['educacao', 'education', 'educacion', 'education']:
    if key in secoes:
        education_section = secoes[key]
        break

if education_section:
    template.add_section_title(doc, get_section_title(education_section))
    
    # Obter lista de formações
    degrees = []
    for key in ['formacao', 'degrees', 'formacion', 'diplomes']:
        if key in education_section:
            degrees = education_section[key]
            break
    
    for degree in degrees:
        doc.add_paragraph(degree)

# Em Andamento - procurar por várias possíveis chaves
in_progress_section = None
for key in ['emAndamento', 'inProgress', 'enProgreso', 'enCours']:
    if key in secoes:
        in_progress_section = secoes[key]
        break

if in_progress_section:
    template.add_section_title(doc, get_section_title(in_progress_section))
    
    # Obter lista de cursos
    courses = []
    for key in ['cursos', 'courses', 'cursos', 'cours']:
        if key in in_progress_section:
            courses = in_progress_section[key]
            break
    
    for course in courses:
        doc.add_paragraph(course)

# Nome do arquivo de saída
output_path = get_field(data, 'nomeArquivoSaida', 'outputFileName')
if not output_path:
    # Se nenhum nome de arquivo for especificado, criar um a partir do nome e idioma
    if nome:
        output_path = f"Curriculo_{nome.replace(' ', '_')}_{selected_lang}.docx"
    else:
        output_path = f"Curriculo_{selected_lang}.docx"

doc.save(output_path)

print(f"Arquivo salvo como: {output_path}")
