"""
Default DOCX CV Template Module.

This module provides functions to create and structure a CV in DOCX format
using the python-docx library. It defines how different sections like titles,
summaries, experiences, and skills are visually represented in the document.
This is the base template for DOCX generation.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
# from docx.oxml.parser import OxmlElement # Not used directly in this version
# from docx.oxml.ns import qn # Not used directly in this version
# from docx.enum.table import WD_TABLE_ALIGNMENT # Not used directly in this version
# from docx.oxml import parse_xml # Not used directly in this version

# Default color and font settings (can be customized or extended)
DEFAULT_ACCENT_COLOR = RGBColor(47, 117, 181)  # A shade of blue
DEFAULT_FONT_SIZE_NORMAL = Pt(11)
DEFAULT_FONT_SIZE_SECTION_TITLE = Pt(14)
DEFAULT_FONT_SIZE_MAIN_TITLE = Pt(22)


def create_document() -> Document:
    """
    Creates and returns a new empty `python-docx` Document object.

    Returns:
        Document: A new `docx.Document` instance.
    """
    return Document()


def add_title(
    doc: Document,
    name: Optional[str],
    email: Optional[str],
    phone: Optional[str],
    linkedin: Optional[str]
) -> None:
    """
    Adds the main title section to the DOCX document, including contact information.

    Args:
        doc (Document): The `python-docx` Document object to add the title to.
        name (Optional[str]): The full name of the CV owner.
        email (Optional[str]): The email address.
        phone (Optional[str]): The phone number.
        linkedin (Optional[str]): The LinkedIn profile URL or vanity name.
    """
    if not name:
        name = "Nome Não Especificado" # Fallback

    # Add name, large and bold
    p_name = doc.add_paragraph()
    run_name = p_name.add_run(name)
    run_name.font.size = DEFAULT_FONT_SIZE_MAIN_TITLE
    run_name.font.bold = True
    p_name.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Contact information line
    p_contact = doc.add_paragraph()
    contact_parts = []
    if email:
        contact_parts.append(("📧 ", email))
    if phone:
        contact_parts.append(("📱 ", phone))
    if linkedin:
        # Assuming a simple text link for DOCX, not a hyperlink object
        contact_parts.append(("🌐 ", linkedin))

    for i, (icon, text) in enumerate(contact_parts):
        if i > 0:
            p_contact.add_run("   ") # Separator
        run_icon = p_contact.add_run(icon)
        run_icon.font.bold = True # Optional: make icons bold
        p_contact.add_run(text)
    p_contact.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Horizontal line separator
    # Using a long dash character. For a graphical line, a shape or border would be needed.
    doc.add_paragraph().add_run("―" * 75) # Adjust length as needed for A4 width


def add_section_title(doc: Document, title: str) -> None:
    """
    Adds a formatted section title to the document.

    The title is preceded by a colored square character for visual emphasis.

    Args:
        doc (Document): The `python-docx` Document object.
        title (str): The text of the section title.
    """
    p_section_title = doc.add_paragraph()

    # Colored square prefix
    run_marker = p_section_title.add_run("■ ")
    run_marker.font.color.rgb = DEFAULT_ACCENT_COLOR
    run_marker.font.size = DEFAULT_FONT_SIZE_SECTION_TITLE
    run_marker.bold = True # Marker is also bold

    # Section title text
    run_title = p_section_title.add_run(title)
    run_title.font.size = DEFAULT_FONT_SIZE_SECTION_TITLE
    run_title.bold = True
    p_section_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


def add_skill_bar(
    doc: Document,
    skill_name: str,
    level: int,
    max_level: int = 5
) -> None:
    """
    Adds a skill with a visual representation of proficiency (skill bar).

    Args:
        doc (Document): The `python-docx` Document object.
        skill_name (str): The name of the skill.
        level (int): The proficiency level of the skill.
        max_level (int): The maximum possible proficiency level. Defaults to 5.
    """
    if not isinstance(level, int) or not isinstance(max_level, int) or level < 0 or max_level <= 0:
        level = 0 # Default to 0 if invalid
        max_level = 5

    level = min(level, max_level) # Ensure level doesn't exceed max_level

    p_skill = doc.add_paragraph()
    run_skill_name = p_skill.add_run(f"{skill_name}: ")
    run_skill_name.font.size = DEFAULT_FONT_SIZE_NORMAL

    # Visual representation of the skill bar
    filled_char = "■"  # Character for filled part of the bar
    empty_char = "□"   # Character for empty part of the bar
    bar_representation = (filled_char * level) + (empty_char * (max_level - level))

    run_bar = p_skill.add_run(bar_representation)
    run_bar.font.color.rgb = DEFAULT_ACCENT_COLOR
    run_bar.font.size = DEFAULT_FONT_SIZE_NORMAL # Match skill name font size


def add_page_break(doc: Document) -> None:
    """
    Adds a page break to the document.

    Args:
        doc (Document): The `python-docx` Document object.
    """
    # Adding a paragraph and then a page break run is a common way
    # Ensure it doesn't add extra vertical space if paragraph is not needed
    # Forcing it into a new paragraph if doc.paragraphs is empty or last one isn't empty
    if not doc.paragraphs or doc.paragraphs[-1].text:
        para_for_break = doc.add_paragraph()
        para_for_break.add_run().add_break(WD_BREAK.PAGE)
    else: # If last paragraph is empty, use it
        doc.paragraphs[-1].add_run().add_break(WD_BREAK.PAGE)


if __name__ == '__main__':
    # Example usage: Create a dummy DOCX file using this template's functions
    print("--- Testing template_docx.py ---")
    dummy_doc = create_document()
    add_title(dummy_doc, "Seu Nome Completo", "email@example.com", "(11) 98765-4321", "linkedin.com/in/seunome")
    
    add_section_title(dummy_doc, "Resumo Profissional")
    dummy_doc.add_paragraph("Este é um parágrafo de exemplo para o resumo profissional...")
    
    add_section_title(dummy_doc, "Habilidades")
    add_skill_bar(dummy_doc, "Python", 4)
    add_skill_bar(dummy_doc, "Comunicação", 5)
    add_skill_bar(dummy_doc, "Design Gráfico", 2, max_level=7) # Example with different max_level

    add_page_break(dummy_doc)
    add_section_title(dummy_doc, "Experiência")
    dummy_doc.add_paragraph("Outra seção de exemplo após uma quebra de página.")

    output_filename = "_test_template_docx_output.docx"
    try:
        dummy_doc.save(output_filename)
        print(f"Documento de teste salvo como: {output_filename}")
        # Check if file exists
        assert os.path.exists(output_filename), "Test DOCX file was not created."
        print("Test DOCX file created successfully.")
        # os.remove(output_filename) # Clean up
        # print(f"Test file {output_filename} removed.")
    except Exception as e:
        print(f"Erro ao salvar documento de teste: {e}")
    
    print("--- template_docx.py test finished ---")
