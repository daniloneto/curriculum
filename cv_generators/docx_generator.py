"""
DOCX CV Generator.

This module provides the DocxGenerator class, which is responsible for
creating CVs in the DOCX format using data loaded by DataLoader and a
specific DOCX template module.
"""
import os
from docx import Document  # For creating and manipulating DOCX files
from typing import List, Dict, Any, Callable, Optional, TYPE_CHECKING

from .base_generator import BaseCvGenerator

# Ensure parent directory is in path for utils import.
# This is a common pattern for modules within packages.
import sys
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from utils import get_field # Used for accessing specific fields in nested job dictionaries

if TYPE_CHECKING:
    from data_loader import DataLoader # For type hinting
    # Assuming template_module is a module object with specific methods.
    # Using 'Any' for now, can be replaced by a Protocol if template methods are standardized.
    from typing import Any as TemplateModuleType


class DocxGenerator(BaseCvGenerator):
    """
    Generates CVs in DOCX format.

    Inherits from BaseCvGenerator and implements the DOCX-specific
    generation logic, utilizing a template module for styling and
    structural components.

    Attributes:
        doc (Optional[Document]): The python-docx Document object.
                                  Initialized in `_setup_document`.
    """

    def __init__(self, data_loader: 'DataLoader', template_module: 'TemplateModuleType'):
        """
        Initializes the DocxGenerator.

        Args:
            data_loader (DataLoader): An instance of DataLoader providing CV data.
            template_module (TemplateModuleType): The loaded template module for DOCX,
                                                  expected to have methods like `create_document`,
                                                  `add_title`, `add_section_title`, etc.
        """
        super().__init__(data_loader, template_module)
        self.expected_extension: str = ".docx"
        self.doc: Optional[Document] = None

    def _setup_document(self) -> None:
        """
        Initializes a new DOCX document using the provided template module.

        Raises:
            TypeError: If the template module's `create_document()` method
                       does not return a `docx.Document` instance.
        """
        self.doc = self.template.create_document()
        if not isinstance(self.doc, Document):
            raise TypeError(
                "Template's create_document() did not return a docx.Document object."
            )

    def _add_personal_info(self) -> None:
        """Adds personal information (name, email, phone, LinkedIn) to the document."""
        if not self.doc:
            raise Exception("Document not initialized. Call _setup_document first.")

        personal_info = self.loader.get_personal_info()
        self.template.add_title(
            self.doc,
            personal_info.get('name'),
            personal_info.get('email'),
            personal_info.get('phone'),
            personal_info.get('linkedin')
        )

    def _add_section(
        self,
        section_key_variants: List[str],
        content_handler_func: Callable[[Dict[str, Any]], None],
        *handler_args: Any
    ) -> bool:
        """
        Adds a generic section to the document if data for it exists.

        Args:
            section_key_variants (List[str]): A list of possible keys for the section
                                              in the CV data.
            content_handler_func (Callable): The method to call to render this
                                             section's content (e.g., _handle_resume_section).
            *handler_args: Additional arguments to pass to the content_handler_func.


        Returns:
            bool: True if the section was found and added, False otherwise.
        """
        if not self.doc:
            raise Exception("Document not initialized. Call _setup_document first.")

        section_data = self.loader.get_section(section_key_variants)
        if section_data:
            self.template.add_section_title(
                self.doc, self.loader.get_section_title(section_data)
            )
            content_handler_func(section_data, *handler_args)
            return True
        return False

    def _handle_resume_section(self, section_data: Dict[str, Any]) -> None:
        """Handles the content for the resume/summary section."""
        if not self.doc: return
        content = self.loader.get_section_content(section_data)
        self.doc.add_paragraph(content)

    def _handle_experience_section(self, section_data: Dict[str, Any]) -> None:
        """Handles the content for the professional experience section."""
        if not self.doc: return
        jobs = self.loader.get_jobs(section_data)
        for job in jobs:
            position = get_field(job, 'cargo', 'position')
            if position:
                self.doc.add_paragraph(str(position), style='List Bullet')

            period = get_field(job, 'periodo', 'period')
            if period:
                self.doc.add_paragraph(str(period))

            description_items: List[str] = []
            # Try to get description using common keys
            raw_desc = get_field(job, 'descricao', 'description', ['descripcion'])
            if isinstance(raw_desc, str):
                description_items = [item.strip() for item in raw_desc.split('\n') if item.strip()]
            elif isinstance(raw_desc, list):
                description_items = [str(item).strip() for item in raw_desc if str(item).strip()]

            for item_text in description_items:
                self.doc.add_paragraph(f"- {item_text}")

    def _handle_skills_section(self, section_data: Dict[str, Any]) -> None:
        """Handles the content for the skills section."""
        if not self.doc: return
        skills = self.loader.get_skills(section_data) # Expects list of dicts
        for skill in skills:
            skill_name = get_field(skill, 'nome', 'name')
            skill_level = get_field(skill, 'nivel', 'level')
            if skill_name and skill_level:
                # Assumes template has add_skill_bar method
                self.template.add_skill_bar(self.doc, str(skill_name), str(skill_level))

    def _handle_list_section(
        self,
        section_data: Dict[str, Any],
        list_keys: List[str],
        item_prefix: Optional[str] = None
    ) -> None:
        """
        Handles content for generic list-based sections (e.g., education, certifications).
        """
        if not self.doc: return
        items = self.loader.get_section_list(section_data, list_keys=list_keys)
        for item in items:
            if item_prefix:
                p = self.doc.add_paragraph()
                p.add_run(item_prefix).bold = True
                p.add_run(str(item))
            else:
                self.doc.add_paragraph(str(item))

    def generate_cv(self) -> str:
        """
        Generates the DOCX CV document.

        It sets up the document, adds personal information, iterates through
        various predefined sections (summary, experience, skills, etc.),
        adds their content, and finally saves the document to the
        `generated_cvs` directory.

        Returns:
            str: The absolute path to the generated DOCX file.
        """
        self._setup_document()
        self._add_personal_info()

        self._add_section(
            ['resumoProfissional', 'professionalSummary', 'resumenProfesional', 'resumo'],
            self._handle_resume_section
        )
        self._add_section(
            ['experienciaProfissional', 'workExperience', 'experienciaLaboral'],
            self._handle_experience_section
        )

        if hasattr(self.template, 'add_page_break'):
            self.template.add_page_break(self.doc)

        self._add_section(
            ['habilidadesTecnicas', 'technicalSkills', 'competencesTechniques', 'habilidades'],
            self._handle_skills_section
        )
        self._add_section(
            ['certificacoes', 'certifications', 'certificaciones'],
            self._handle_list_section, ['lista', 'list', 'items'], "🏅 "
        )
        self._add_section(
            ['educacao', 'education', 'educacion'],
            self._handle_list_section, ['formacao', 'degrees', 'formacion', 'diplomes', 'items']
        )
        self._add_section(
            ['emAndamento', 'inProgress', 'enProgreso', 'enCours'],
            self._handle_list_section, ['cursos', 'courses', 'items']
        )

        absolute_output_path = self._get_full_output_path(
            default_prefix="Curriculo",
            extension=self.expected_extension
        )

        if not self.doc:
             raise Exception("Document not initialized before save. Call _setup_document first.")
        self.doc.save(absolute_output_path)
        return absolute_output_path


if __name__ == '__main__':
    # This block is for illustrative purposes.
    # Direct execution requires setting up DataLoader, a DOCX template module,
    # and ensuring utils.py and data_loader.py are accessible in the Python path.
    print(
        "DocxGenerator class defined. For testing, run the main CLI script "
        "(e.g., curriculo_docx.py) or create dedicated unit/integration tests."
    )

    # Example of a more complete direct test (requires careful setup):
    # import json
    # from data_loader import DataLoader # Adjust path if necessary
    # from templates import TemplateManager # Adjust path if necessary
    #
    # print("\n--- Running DocxGenerator Direct Test Example ---")
    # # Ensure the test output base directory exists
    # test_output_base_dir = os.path.abspath(
    #    os.path.join(os.path.dirname(__file__), '..', '..', '_test_direct_run_output')
    # )
    # os.makedirs(test_output_base_dir, exist_ok=True)
    #
    # # Create a dummy JSON file for testing in the test output base directory
    # dummy_json_filename = "curriculo_test_direct_docx.json"
    # dummy_json_path = os.path.join(test_output_base_dir, dummy_json_filename)
    # with open(dummy_json_path, "w", encoding="utf-8") as f:
    #     json.dump({
    #         "languageName": "TestishDirect",
    #         "nome": "Direct Doc Tester", "email": "direct_doc@test.com",
    #         "secoes": {
    #             "resumo": {"titulo": "Direct Test Summary", "texto": "This is a direct test."}
    #         }
    #     }, f)
    #
    # try:
    #     # DataLoader root_dir should be where curriculo_*.json files are,
    #     # or where the custom JSON file is (if json_file_path is relative).
    #     # For this test, json_file_path is absolute.
    #     loader = DataLoader(
    #         language_code="test_direct_docx",
    #         json_file_path=dummy_json_path, # Provide absolute path
    #         root_dir=test_output_base_dir # This is where generated_cvs will be created
    #     )
    #
    #     # TemplateManager needs to find template modules.
    #     # Assuming 'templates' dir is sibling to 'cv_generators' and 'data_loader'
    #     project_true_root = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
    #     tm = TemplateManager(root_dir=project_true_root)
    #     docx_template_module = tm.get_template('docx')
    #
    #     generator = DocxGenerator(data_loader=loader, template_module=docx_template_module)
    #     generated_file = generator.generate_cv()
    #     print(f"Generated test DOCX at: {generated_file}")
    #
    #     assert os.path.exists(generated_file), "Generated DOCX file not found."
    #     assert "generated_cvs" in generated_file, "File not in generated_cvs subdir."
    #     print("Direct test for DocxGenerator seems successful.")
    #
    # except Exception as e_test:
    #     print(f"Error during DocxGenerator direct test: {e_test}")
    #     import traceback
    #     traceback.print_exc()
    #
    # finally:
    #     # Clean up the dummy JSON file
    #     if os.path.exists(dummy_json_path):
    #         os.remove(dummy_json_path)
    #     # Consider cleaning up generated_cvs for this specific test run if needed
    #     # For example, by removing the specific file or the "generated_cvs" dir
    #     # if it was created solely for this test in test_output_base_dir.
    #     generated_cvs_path = os.path.join(test_output_base_dir, "generated_cvs")
    #     if os.path.exists(generated_cvs_path) and not os.listdir(generated_cvs_path):
    #         os.rmdir(generated_cvs_path)
    #     if os.path.exists(test_output_base_dir) and not os.listdir(test_output_base_dir):
    #          os.rmdir(test_output_base_dir)
    # print("--- DocxGenerator Direct Test Example Finished ---")
